from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import csv

def set_cell_margins(cell, top, bottom, left, right):
    """Set custom margins for a cell in a Word document."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, value in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        margin_el = OxmlElement(f'w:{margin}')
        margin_el.set(qn('w:w'), str(value))
        margin_el.set(qn('w:type'), "dxa")  # dxa is twentieths of a point
        tcMar.append(margin_el)
    tcPr.append(tcMar)

def create_nested_table(cell, rows, cols, items, prices, old_prices, font_name, font_size_items, font_size_main, font_size_superscript, font_bold_items, font_bold):
    nested_table = cell.add_table(rows=rows, cols=2)  # Ensure the table has 3 cols as needed for item, price, and old price

    for i, row in enumerate(nested_table.rows):
        item_cell = row.cells[0]
        price_cell = row.cells[1]

        # Set padding for the first and last cells
        set_cell_margins(item_cell, 0, 0, 800, 0)  # 200 Twips = 0.2 Inches left padding for the first cell
        set_cell_margins(price_cell, 0, 0, 0, 800)  # 200 Twips = 0.2 Inches right padding for the last cell

        # Existing formatting code
        item_paragraph = item_cell.paragraphs[0]
        item_run = item_paragraph.add_run(items[i] if i < len(items) else "")
        item_run.font.name = font_name
        item_run.font.size = font_size_items
        item_run.bold = font_bold_items

        def format_price_cell(price_cell, price, strike=False):
            price_paragraph = price_cell.paragraphs[0]
            price_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            price_parts = price.split('.')
            main_price = price_parts[0] + "."
            cents = price_parts[1] if len(price_parts) > 1 else ""
            run_main = price_paragraph.add_run(main_price)
            run_main.font.name = font_name
            run_main.font.size = font_size_main
            run_main.bold = font_bold
            run_main.font.strike = strike

            if cents:
                run_superscript = price_paragraph.add_run(cents)
                run_superscript.font.name = font_name
                run_superscript.font.size = font_size_superscript
                run_superscript.bold = font_bold
                run_superscript.font.superscript = True
                run_superscript.font.strike = strike

        format_price_cell(price_cell, prices[i] if i < len(prices) else "")

    return nested_table
def add_rectangle(items, prices, old_prices):
    rect_table = doc.add_table(rows=1, cols=1)
    rect_table.style = 'Table Grid'
    rect_cell = rect_table.cell(0, 0)
    
    # Adjusted to increase the width by about 2 inches (approximately 5 cm)
    rect_cell.width = Inches(width)
    rect_cell.height = Inches(height + 0.12)  # Set the height of the rectangle as before
    rect_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adjust row height of the rectangle
    tr = rect_table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int((height + 0.12) * 1440)))  # Convert height to Word's unit
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)

    create_nested_table(rect_cell, len(items), 3, items, prices, old_prices, font_name, font_size_items, font_size_main, font_size_superscript, font_bold_items, font_bold)


# Create a new Document
try:
    doc = Document()
except Exception as e:
    with open ('logfile.log', 'a') as logfile:
        logfile.write('Problem z stworzeniem dokumentu\n')
        logfile.write(str(e))
        logfile.write("\n")
    

# Adjust document margins
try:
    for section in doc.sections:
        section.top_margin = Inches(0.1)  # Top margin set to 0.5 inch
        section.bottom_margin = Inches(0.1)  # Bottom margin set to 0.5 inch
        section.left_margin = Inches(0.1)  # Left margin set to 0.5 inch
        section.right_margin = Inches(0.1)  # Right margin set to 0.5 inch
except Exception as e:
    with open ('logfile.log', 'a') as logfile:
        logfile.write('Problem z z dodaniem marginesów\n')
        logfile.write(str(e))
        logfile.write("\n")

# Define a font style for the items, the main part of the price, and the superscript part of the price
try:
    font_name = "Montserrat"
    font_size_items = Pt(32)  # Adjust the size as needed for item names
    font_size_main = Pt(24)
    font_size_superscript = Pt(24)  # Smaller font size for the superscript
    font_bold_items = True  # Set to True if item names should be bold
    font_bold = True  # Set to True if prices should be bold
except Exception as e:
    with open ('logfile.log', 'a') as logfile:
        logfile.write('Problem z wczytaniem ustawien\n')
        logfile.write(str(e))
        logfile.write("\n")
# Sample data for the table
items1 = []
items2 = []
prices1 = []
prices2 = []
old_prices1 = []
old_prices2 = []

try:
    with open('ceny regularne.csv', encoding='utf-8', newline='') as csvfile:
        csv_reader = csv.reader(csvfile, delimiter=',')
        next(csv_reader)
        for row in csv_reader:
            if row[2] == '1':
                items1.append(row[0])
                prices1.append(row[1])
            elif row[2] == '2':
                items2.append(row[0])
                prices2.append(row[1])
except Exception as e:
    with open ('logfile.log', 'a') as logfile:
        logfile.write('Problem z otwarciem pliku ceny regularne\n')
        logfile.write(str(e))
        logfile.write("\n")
                
            
# mid_index = len(items) // 2 + len(items) % 2 
# items1, items2 = items[:mid_index], items[mid_index:]
# prices1, prices2 = prices[:mid_index], prices[mid_index:]
# old_prices1, old_prices2 = old_prices[:mid_index], old_prices[mid_index:]

width = 7.09 
height = 4.9  
try:
    add_rectangle(items1, prices1, old_prices1)
    add_rectangle(items2, prices2, old_prices2)
except Exception as e:
    with open ('logfile.log', 'a') as logfile:
        logfile.write('Problem z dodaniem prostokątów\n')
        logfile.write(str(e))
        logfile.write("\n")
        
        
import os
# Save the document to a file
try:
    doc_path ="Etykieta do druku ceny regularne.docx"
    #doc.save(doc_path)
    # Saving the document as DOCX
    doc.save(doc_path)
    # Converting the DOCX to PDF

except Exception as e:
    with open ('logfilesale.log', 'a') as logfile:
        logfile.write('Problem z zapisem pliku\n')
        print(e)
        logfile.write(str(e))
        logfile.write("\n")
