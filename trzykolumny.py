from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import csv

def create_nested_table(cell, rows, cols, items, prices, old_prices, font_name, font_size_items, font_size_main, font_size_superscript, font_bold_items, font_bold):
    # Create a table in the given cell with an additional column for old prices
    nested_table = cell.add_table(rows=rows, cols=3)

    # Set the style and format for the nested table
    for i, row in enumerate(nested_table.rows):
        item_cell = row.cells[0]
        price_cell_1 = row.cells[1]
        price_cell_2 = row.cells[2]  # The new third column for old prices

        # Set item name and format
        item_paragraph = item_cell.paragraphs[0]
        item_run = item_paragraph.add_run(items[i] if i < len(items) else "")
        item_run.font.name = font_name
        item_run.font.size = font_size_items
        item_run.bold = font_bold_items

        # Function to format price cells, with an option to strike through the price
        def format_price_cell(price_cell, price, strike=False):
            price_paragraph = price_cell.paragraphs[0]
            price_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Splitting price into main and cents
            price_parts = price.split('.')
            main_price = price_parts[0] + "." if len(price_parts) > 0 else ""
            cents = price_parts[1] if len(price_parts) > 1 else ""

            # Add run for the main part of the price
            run_main = price_paragraph.add_run(main_price)
            run_main.font.name = font_name
            run_main.font.size = font_size_main
            run_main.bold = font_bold
            run_main.font.strike = strike

            # Add run for the superscript part of the price, if applicable
            if cents:
                run_superscript = price_paragraph.add_run(cents)
                run_superscript.font.name = font_name
                run_superscript.font.size = font_size_superscript
                run_superscript.bold = font_bold
                run_superscript.font.superscript = True
                run_superscript.font.strike = strike

        # Format and populate first price cell
        format_price_cell(price_cell_1, prices[i] if i < len(prices) else "")

        # Format and populate second price cell with old prices, applying strikethrough
        format_price_cell(price_cell_2, old_prices[i] if i < len(old_prices) else "", strike=True)
    
    return nested_table

# Create a new Document
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.5)  # Top margin set to 0.5 inch
    section.bottom_margin = Inches(0.5)  # Bottom margin set to 0.5 inch
    section.left_margin = Inches(0.5)  # Left margin set to 0.5 inch
    section.right_margin = Inches(0.5)  # Right margin set to 0.5 inch


# Define a font style for the items, the main part of the price, and the superscript part of the price
font_name = "Montserrat"
font_size_items = Pt(32)  # Adjust the size as needed for item names
font_size_main = Pt(24)
font_size_superscript = Pt(24)  # Smaller font size for the superscript
font_bold_items = True  # Set to True if item names should be bold
font_bold = True  # Set to True if prices should be bold

# Sample data for the table
items1 = []
items2 = []
prices1 = []
prices2 = []
old_prices1 = []
old_prices2 = []

with open('dane_3_kolumny.csv', encoding='utf-8', newline='') as csvfile:
    csv_reader = csv.reader(csvfile, delimiter=';')
    next(csv_reader)
    for row in csv_reader:
        if row[3] == '1':
            items1.append(row[0])
            prices1.append(row[1])
            old_prices1.append(row[2])
        elif row[3] == '2':
            items2.append(row[0])
            prices2.append(row[1])
            old_prices2.append(row[2])
        

# Divide your items, prices, and old_prices into two groups
# mid_index = len(items) // 2 + len(items) % 2  # Adjust to ensure at least one item in each if uneven
# items1, items2 = items[:mid_index], items[mid_index:]
# prices1, prices2 = prices[:mid_index], prices[mid_index:]
# old_prices1, old_prices2 = old_prices[:mid_index], old_prices[mid_index:]

# Define the width and height for the rectangular box, adjusted to fit two on one page
width = 7.09  # width in inches
height = 4.9  # height in inches, adjusted for two rectangles

# Function to add rectangles
def add_rectangle(items, prices, old_prices):
    rect_table = doc.add_table(rows=1, cols=1)
    rect_table.style = 'Table Grid'
    rect_cell = rect_table.cell(0, 0)
    rect_cell.width = Inches(width)
    rect_cell.height = Inches(height)  # Set the height of the rectangle
    rect_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adjust row height of the rectangle
    tr = rect_table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height * 1440)))  # Convert height to Word's unit
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)

    create_nested_table(rect_cell, len(items), 3, items, prices, old_prices, font_name, font_size_items, font_size_main, font_size_superscript, font_bold_items, font_bold)

# Add the first rectangle with the first set of items
add_rectangle(items1, prices1, old_prices1)

# Add the second rectangle with the second set of items
add_rectangle(items2, prices2, old_prices2)

# Save the document to a file
doc.save("trzy_kolumny.docx")

print("Styled document with table created successfully.")
